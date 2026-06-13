"""
真实简历解析引擎
从 PDF/DOCX/TXT 中提取文本，然后用关键词+正则提取：
- 技能列表
- 教育背景
- 工作年限
- 意向岗位
- 意向行业
"""

import re
import os
import json
from pathlib import Path
from typing import Optional

# ── 技能词典（中英文覆盖，约200+技能）─────────────────
SKILL_DICT = {
    # 编程语言
    "python", "java", "javascript", "typescript", "go", "golang", "rust", "c++", "c#", "c语言",
    "php", "ruby", "swift", "kotlin", "scala", "shell", "bash", "matlab", "r语言",
    # 数据库
    "sql", "mysql", "postgresql", "mongodb", "redis", "oracle", "sqlserver", "elasticsearch",
    "hbase", "cassandra", "neo4j", "influxdb", "clickhouse",
    # 大数据
    "hadoop", "spark", "hive", "flink", "kafka", "storm", "airflow", "presto",
    "数据仓库", "数据湖", "etl", "数据建模",
    # 机器学习/AI
    "机器学习", "深度学习", "自然语言处理", "计算机视觉", "nlp", "cv",
    "推荐系统", "强化学习", "知识图谱", "图神经网络",
    "pytorch", "tensorflow", "keras", "scikit-learn", "xgboost", "lightgbm",
    "pandas", "numpy", "scipy", "opencv", "huggingface", "langchain",
    # 数据分析
    "数据分析", "数据挖掘", "数据可视化", "a/b测试", "统计学", "回归分析",
    "excel", "tableau", "power bi", "finebi", "metabase", "superset",
    # 前端
    "react", "vue", "angular", "svelte", "next.js", "nuxt",
    "html5", "css3", "sass", "less", "webpack", "vite", "babel",
    "微信小程序", "uni-app", "flutter", "react native", "electron",
    # 后端
    "node.js", "express", "nestjs", "django", "flask", "fastapi",
    "spring", "spring boot", "spring cloud", "mybatis", "hibernate",
    "微服务", "restful api", "graphql", "grpc", "消息队列",
    # DevOps/云
    "docker", "kubernetes", "k8s", "jenkins", "gitlab ci", "github actions",
    "aws", "azure", "gcp", "阿里云", "腾讯云", "华为云",
    "nginx", "linux", "terraform", "ansible", "prometheus", "grafana",
    # 测试
    "自动化测试", "selenium", "jmeter", "pytest", "junit", "接口测试",
    "性能测试", "压力测试", "单元测试", "集成测试", "ci/cd",
    # 产品/设计
    "产品设计", "产品规划", "需求分析", "用户研究", "竞品分析",
    "prd", "mrd", "roadmap", "用户画像", "用户增长",
    "axure", "figma", "sketch", "photoshop", "illustrator", "adobe xd",
    "ui设计", "ux设计", "交互设计", "视觉设计", "动效设计",
    "blender", "cinema 4d", "after effects", "premiere",
    # 项目管理
    "项目管理", "敏捷开发", "scrum", "kanban", "pmp",
    "jira", "confluence", "teambition", "飞书", "notion",
    # 运营
    "新媒体运营", "内容运营", "社群运营", "电商运营", "活动运营",
    "用户运营", "seo", "sem", "信息流投放", "kol运营",
    "私域运营", "增长黑客", "数据运营", "产品运营",
    # 市场
    "市场营销", "品牌管理", "品牌策划", "公关", "广告投放",
    "整合营销", "事件营销", "社交媒体营销", "内容营销",
    # 人力资源
    "人力资源", "招聘", "薪酬", "绩效", "培训", "hrbp",
    "员工关系", "组织发展", "人才发展", "企业文化",
    "六模块", "三支柱", "hris",
    # 财务
    "财务管理", "审计", "税务", "预算管理", "成本控制",
    "财务报表", "财务分析", "会计准则", "cpa",
    # 供应链
    "供应链管理", "物流管理", "采购管理", "仓储管理", "质量管理",
    # 法务
    "法务", "合规", "合同管理", "知识产权", "公司法",
    # 销售
    "大客户销售", "渠道销售", "商务谈判", "客户关系管理", "crm",
    "销售管理", "销售运营", "商务拓展",
    # 通用技能
    "英语流利", "日语", "韩语", "沟通能力", "团队协作",
    "领导力", "问题解决", "批判性思维", "时间管理",
    "演讲能力", "写作能力", "数据分析能力", "逻辑思维",
    # 安全
    "网络安全", "渗透测试", "安全审计", "安全架构", "零信任",
    # 区块链
    "区块链", "智能合约", "solidity", "web3", "defi",
    # 游戏
    "unity", "unreal", "unreal engine", "游戏策划", "游戏运营",
}

# ── 教育学历正则模式 ───────────────────────────────────
EDUCATION_PATTERNS = [
    (r"(博士|博士研究生|ph\.?d|博士在读|博后)", "博士"),
    (r"(硕士|硕士研究生|master|mba|硕士在读|研究生)", "硕士"),
    (r"(本科|学士|bachelor|本科在读|大学本科)", "本科"),
    (r"(大专|专科|高职)", "大专"),
]

# ── 工作经验年限模式 ───────────────────────────────────
EXPERIENCE_PATTERNS = [
    (r"(\d+)\s*[-~至到]\s*(\d+)\s*年.*?经验", lambda m: f"{m.group(1)}-{m.group(2)}年经验"),
    (r"(\d{1,2})\s*年.*?(工作经验|工作经历|从业|相关经验)", lambda m: f"{m.group(1)}年经验"),
    (r"(应届|毕业|实习生|intern|实习)", "应届/实习"),
    (r"(\d+)\+?\s*年工作经验", lambda m: f"{m.group(1)}+年经验"),
]

# ── 意向岗位关键词 ─────────────────────────────────────
INTENDED_ROLE_KEYWORDS = [
    "数据分析", "数据挖掘", "数据科学", "商业分析", "数据运营",
    "产品经理", "产品总监", "产品运营", "产品设计",
    "后端开发", "前端开发", "全栈开发", "算法工程师", "机器学习",
    "人工智能", "深度学习", "自然语言处理", "计算机视觉",
    "测试工程师", "运维工程师", "安全工程师", "架构师",
    "人力资源", "hrbp", "招聘", "薪酬", "培训", "组织发展",
    "财务管理", "会计", "审计", "财务分析",
    "市场营销", "品牌管理", "新媒体运营", "内容运营",
    "视觉设计", "ui设计", "ux设计", "交互设计",
    "项目经理", "技术经理", "技术总监", "cto",
    "运营", "销售", "法务", "行政", "采购", "供应链",
    "教师", "讲师", "课程设计", "咨询顾问",
]

# ── 意向行业关键词 ─────────────────────────────────────
INDUSTRY_KEYWORDS = [
    ("互联网", ["互联网", "电商", "短视频", "直播", "社交媒体"]),
    ("人工智能", ["人工智能", "ai", "机器学习", "深度学习"]),
    ("金融科技", ["金融科技", "fintech", "支付", "区块链", "web3"]),
    ("金融", ["金融", "银行", "证券", "基金", "保险", "投资"]),
    ("制造业", ["制造", "汽车", "新能源", "芯片", "半导体", "电子"]),
    ("医疗健康", ["医疗", "医药", "健康", "生物", "基因"]),
    ("教育", ["教育", "培训", "在线教育", "k12", "职业教育"]),
    ("快消", ["快消", "零售", "消费品", "食品饮料"]),
    ("咨询", ["咨询", "战略", "管理咨询"]),
    ("游戏", ["游戏", "电竞", "二次元", "元宇宙"]),
    ("房地产", ["房地产", "物业", "建筑"]),
    ("物流", ["物流", "快递", "仓储", "运输"]),
    ("能源", ["能源", "电力", "石油", "新能源"]),
    ("通信", ["通信", "5g", "运营商", "电信"]),
]


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """从 PDF/DOCX/TXT 字节数据中提取文本"""
    ext = Path(filename).suffix.lower() if filename else ""

    if ext in (".txt", ".md", ".csv"):
        return content.decode("utf-8", errors="ignore")

    elif ext == ".pdf":
        try:
            from PyPDF2 import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages_text.append(t)
            return "\n".join(pages_text)
        except Exception as e:
            print(f"[ResumeParser] PDF解析失败: {e}")
            return ""

    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也读表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"[ResumeParser] DOCX解析失败: {e}")
            return ""

    else:
        # 尝试当作文本解码
        try:
            return content.decode("utf-8", errors="ignore")
        except:
            return f"[不支持的文件格式: {ext}]"


def extract_skills(text: str) -> list[str]:
    """从文本中提取技能关键词"""
    found = set()
    text_lower = text.lower()
    for skill in SKILL_DICT:
        if skill.lower() in text_lower:
            found.add(skill)
    # 按在原文中出现顺序排序
    result = sorted(found, key=lambda s: text_lower.index(s.lower()))
    return result


def extract_education(text: str) -> str:
    """提取最高学历"""
    best_rank = {"博士": 5, "硕士": 4, "本科": 3, "大专": 2}
    best = "本科"  # 默认
    best_score = 3
    for pattern, label in EDUCATION_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            score = best_rank.get(label, 0)
            if score > best_score:
                best = label
                best_score = score
    return best


def extract_experience(text: str) -> str:
    """提取工作年限"""
    for pattern, handler in EXPERIENCE_PATTERNS:
        match = re.search(pattern, text)
        if match:
            if callable(handler):
                return handler(match)
            return handler
    return "1-3年经验"


def extract_intended_roles(text: str) -> list[str]:
    """从简历中提取意向岗位"""
    # 方法1：找"求职意向"或"期望岗位"段落
    target_area = ""
    for pattern in [
        r"求职意向[：:](.*?)(?:\n|$)",
        r"期望岗位[：:](.*?)(?:\n|$)",
        r"意向岗位[：:](.*?)(?:\n|$)",
        r"应聘岗位[：:](.*?)(?:\n|$)",
        r"目标岗位[：:](.*?)(?:\n|$)",
        r"期望职位[：:](.*?)(?:\n|$)",
    ]:
        match = re.search(pattern, text)
        if match:
            target_area = match.group(1).strip().lower()
            break

    # 方法2：从文本中提取
    roles = []
    for rk in INTENDED_ROLE_KEYWORDS:
        if rk.lower() in text.lower():
            roles.append(rk)
    # 求职意向段落中的词优先
    roles.sort(key=lambda r: (r.lower() not in target_area, text.lower().index(r.lower())))
    return roles[:4] if roles else ["数据分析"]


def extract_intended_industries(text: str) -> list[str]:
    """从简历中提取意向行业"""
    found = []
    text_lower = text.lower()
    for industry, keywords in INDUSTRY_KEYWORDS:
        if any(kw.lower() in text_lower for kw in keywords):
            found.append(industry)
    return found[:3] if found else ["互联网"]


def extract_education_major(text: str) -> str:
    """提取专业信息"""
    major_patterns = [
        r"(计算机科学|软件工程|数据科学|人工智能|信息管理)",
        r"(统计学|数学|应用数学|计算数学)",
        r"(金融|经济|会计|财务|工商管理)",
        r"(人力资源管理|市场营销|电子商务)",
        r"(电子工程|通信工程|自动化|机械)",
        r"(设计|艺术|传媒|新闻|中文)",
        r"(生物|化学|物理|材料|环境)",
    ]
    for pattern in major_patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)
    return "计算机科学"


def extract_projects(text: str) -> list[str]:
    """提取项目经验名称"""
    # 匹配"项目经验"或"项目经历"段落中的项目名
    projects = []
    # 尝试找项目名称模式
    proj_section = ""
    section_match = re.search(
        r"(项目[经历经验].*?)(?:(?:工作[经历经验])|(?:实习[经历经验])|(?:教育[背景经历])|(?:专业技能)|$)",
        text, re.DOTALL | re.IGNORECASE
    )
    if section_match:
        proj_section = section_match.group(1)
        # 提取项目名称（行首或 • 开头的行）
        lines = proj_section.split("\n")
        for line in lines:
            line = line.strip()
            if len(line) > 4 and len(line) < 60 and not any(
                kw in line for kw in ["项目经历", "项目经验", "负责", "使用", "技术栈"]
            ):
                if re.match(r"^[•·\-\*\d\.、\s]*([\u4e00-\u9fff\w]{2,30})$", line):
                    name = re.sub(r"^[•·\-\*\d\.、\s]+", "", line).strip()
                    if name:
                        projects.append(name)

    if not projects:
        # 从全文中提取典型项目名模式
        proj_patterns = [
            r"([\u4e00-\u9fff]{3,8}(?:系统|平台|项目|模型|分析|优化|设计|开发|搭建))",
        ]
        for pat in proj_patterns:
            matches = re.findall(pat, text)
            projects.extend(matches)

    return list(dict.fromkeys(projects))[:4] if projects else ["数据建模分析", "业务系统优化"]


def extract_work_style(text: str) -> list[str]:
    """提取工作风格偏好"""
    style_keywords = {
        "数据驱动": ["数据驱动", "数据导向"],
        "快节奏": ["快节奏", "高效率", "敏捷"],
        "结果导向": ["结果导向", "目标导向", "kpi"],
        "创新": ["创新", "creative", "探索"],
        "协作": ["协作", "团队合作", "collaboration"],
        "扁平化": ["扁平", "开放", "自驱"],
        "稳定": ["稳定", "长期", "深耕"],
    }
    found = []
    text_lower = text.lower()
    for style, keywords in style_keywords.items():
        if any(kw.lower() in text_lower for kw in keywords):
            found.append(style)
    return found if found else ["数据驱动", "结果导向"]


def extract_deal_breakers(text: str) -> list[str]:
    """从简历中提取不可接受项（deal breakers）"""
    breakers = []
    text_lower = text.lower()
    # 检查是否有明确拒绝的表述
    breaker_patterns = [
        (r"(不接受|不考虑|无法接受|拒绝).{0,10}(加班|996|大小周)", "拒绝高强度加班"),
        (r"(不接受|无法).{0,10}(出差|外派)", "不接受频繁出差"),
        (r"(不接受|不考虑).{0,10}(研发|开发|编程|coding)", "拒绝纯开发岗"),
    ]
    for pattern, label in breaker_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            breakers.append(label)
    if not breakers:
        breakers = ["纯开发岗", "24小时on-call", "无明确晋升路径"]
    return breakers


def parse_resume(content: bytes, filename: str) -> dict:
    """
    完整的简历解析流程：
    1. 提取原始文本
    2. 提取三圈画像要素
    3. 返回结构化结果
    """
    # Step 1: 文本提取
    raw_text = extract_text_from_bytes(content, filename)
    if not raw_text:
        raise ValueError(f"无法从文件 {filename} 中提取文字内容，请确认文件格式为 PDF/DOCX/TXT")

    # Step 2: 各项提取
    skills = extract_skills(raw_text)
    education_level = extract_education(raw_text)
    education_major = extract_education_major(raw_text)
    experience = extract_experience(raw_text)
    intended_roles = extract_intended_roles(raw_text)
    intended_industries = extract_intended_industries(raw_text)
    projects = extract_projects(raw_text)
    work_style = extract_work_style(raw_text)
    deal_breakers = extract_deal_breakers(raw_text)

    # Step 3: 构建三圈画像
    interest_profile = {
        "preferred_industries": intended_industries,
        "preferred_roles": intended_roles,
        "work_style": work_style,
    }

    ability_profile = {
        "skills": skills[:12] if len(skills) > 12 else skills,
        "education": f"{education_level} | {education_major}",
        "experience": experience,
        "projects": projects,
    }

    return {
        "interest_profile": interest_profile,
        "ability_profile": ability_profile,
        "deal_breakers": deal_breakers,
        "raw_text_preview": raw_text[:300],
        "extracted_skills": skills,
        "extracted_education": education_level,
        "extracted_major": education_major,
        "extracted_experience": experience,
        "extracted_roles": intended_roles,
        "extracted_industries": intended_industries,
    }
